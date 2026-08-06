import os
import json
import time
import shutil
import pdfplumber
from typing import List, Optional
from pydantic import BaseModel, Field
from google import genai
from google.genai import types
import docx

API_KEY = "YOUR_KEY_HERE"  # Replace with your actual key
client = genai.Client(api_key=API_KEY.strip())

# ==========================================
# MODELS FOR CRITERIA EXTRACTION (multiple roles)
# ==========================================
class RoleCriteria(BaseModel):
    name: str = Field(description="Name of the role (e.g., 'Software Engineer')")
    requirements: List[str] = Field(description="All requirements mentioned for this role")
    degree_level: Optional[str] = Field(description="Required degree level if explicitly mentioned")
    qualification: Optional[str] = Field(description="Required major/qualification if explicitly mentioned")
    experience_years: Optional[float] = Field(description="Minimum years of experience if explicitly mentioned")

class CriteriaExtractionResult(BaseModel):
    roles: List[RoleCriteria] = Field(description="List of roles found in the document")

# ==========================================
# MODELS FOR CV EVALUATION
# ==========================================
class EvaluationResult(BaseModel):
    extracted_degree: str = Field(description="Candidate's highest degree level")
    extracted_qualification: str = Field(description="Candidate's major/field of study")
    extracted_experience: float = Field(description="Candidate's total years of professional experience")
    certainty_score: int = Field(description="Confidence score (0-100) that candidate meets ALL requirements")
    reason: str = Field(description="Short explanation of why the candidate was accepted, doubted, or rejected")

# ==========================================
# EXTRACT CRITERIA FROM RECRUITER DOCX (with table support)
# ==========================================
def extract_criteria_from_docx(docx_path):
    """
    Read .docx and extract both paragraph text and table content.
    Tables are converted to a markdown-like format to preserve column structure.
    """
    try:
        doc = docx.Document(docx_path)
        # Paragraphs
        para_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

        # Tables – extract as markdown tables
        table_text = ""
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
            if rows:
                if len(rows) > 1:
                    header_cells = len(table.rows[0].cells)
                    separator = "| " + " | ".join(["---"] * header_cells) + " |"
                    rows.insert(1, separator)
                table_text += "\n" + "\n".join(rows) + "\n"

        full_text = para_text + "\n\n" + table_text
    except Exception as e:
        print(f"Error reading docx: {e}")
        return []

    if not full_text.strip():
        return []

    prompt = (
        "Extract all job roles and their requirements from the following document. "
        "The document may contain tables that compare requirements for different roles (e.g., Medical Path vs Technical Path). "
        "Treat each distinct path/role as a separate role. For each role, extract the role name, a list of requirements (as separate statements), "
        "and if explicitly mentioned, the required degree level, qualification/major, and minimum years of experience. "
        "If a table has columns for different roles, use the column headers as role names and the cell contents as requirements for that role. "
        "Output JSON with key 'roles' which is a list of objects, each with: "
        "name (string), requirements (list of strings), degree_level (string or null), "
        "qualification (string or null), experience_years (number or null).\n\n"
        f"Document content:\n{full_text}"
    )

    try:
        response = client.models.generate_content(
            model='gemini-3.1-flash-lite',
            contents=prompt,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=CriteriaExtractionResult,
                temperature=0.0
            )
        )
        data = json.loads(response.text)
        roles = data.get("roles", [])
        for role in roles:
            role.setdefault("requirements", [])
            role.setdefault("degree_level", "")
            role.setdefault("qualification", "")
            role.setdefault("experience_years", 0.0)
        return roles
    except Exception as e:
        print(f"Criteria extraction error: {e}")
        return []

# ==========================================
# EVALUATE A SINGLE CV AGAINST ONE ROLE
# ==========================================
def evaluate_cv(resume_text, role: RoleCriteria):
    req_bullets = "\n".join(f"- {r}" for r in role.requirements if r.strip())
    core = []
    if role.degree_level:
        core.append(f"Degree: {role.degree_level}")
    if role.qualification:
        core.append(f"Major: {role.qualification}")
    if role.experience_years and role.experience_years > 0:
        core.append(f"Exp: {role.experience_years} yrs")
    core_bullets = "\n".join(f"- {c}" for c in core) if core else ""
    all_reqs = req_bullets + ("\n" + core_bullets if core_bullets else "")

    # ----- UPDATED PROMPT WITH STRICTER RULES -----
    prompt = (
        "Score 0-100 vs requirements. Rules:\n"
        "- The required degree level and qualification must be exact or very closely related.\n"
        "- Degrees in unrelated fields (e.g., Healthcare Management for Finance) should be penalised heavily (score ≤ 30).\n"
        "- MBBS/MD must be present if required.\n"
        "- Experience years must be met; if not, reduce score proportionally.\n"
        "- Provide reason (max 10 words). If score < 80, list missing criteria.\n"
        "Requirements:\n" + all_reqs + "\n"
        "CV:\n" + resume_text[:3000] + "\n"
        "JSON: extracted_degree, extracted_qualification, extracted_experience (float), certainty_score (int), reason (string)."
    )

    max_retries = 5
    delay = 15
    for attempt in range(max_retries):
        try:
            time.sleep(6)
            response = client.models.generate_content(
                model='gemini-3.1-flash-lite',
                contents=prompt,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    response_schema=EvaluationResult,
                    temperature=0.0
                )
            )
            return json.loads(response.text)
        except Exception as e:
            err = str(e)
            if "429" in err or "RESOURCE_EXHAUSTED" in err.upper():
                print(f"⚠️ Rate limit, retry in {delay}s")
                time.sleep(delay)
                delay *= 2
            else:
                print(f"LLM evaluation error: {e}")
                return {
                    "extracted_degree": "",
                    "extracted_qualification": "",
                    "extracted_experience": 0.0,
                    "certainty_score": 0,
                    "reason": f"Evaluation failed: {str(e)}"
                }
    return {
        "extracted_degree": "",
        "extracted_qualification": "",
        "extracted_experience": 0.0,
        "certainty_score": 0,
        "reason": "Unable to evaluate after retries"
    }

# ==========================================
# PDF TEXT EXTRACTION
# ==========================================
def extract_pdf_text(pdf_path):
    try:
        with pdfplumber.open(pdf_path) as pdf:
            return "\n".join(p.extract_text(layout=True) or "" for p in pdf.pages)
    except Exception as e:
        print(f"PDF error {os.path.basename(pdf_path)}: {e}")
        return ""

# ==========================================
# GENERATOR – Single assignment per CV
# ==========================================
def main_generator(roles: List[RoleCriteria]):
    """
    Yields events for each CV, assigning it to the best‑matched role.
    If no role scores ≥ 40, the CV goes to an 'Other' bucket.
    """
    # Prepare folders per role + an "Other" folder
    all_folders = {}
    for role in roles:
        safe_name = role.name.replace(" ", "_").replace("/", "_")
        base = f"role_{safe_name}"
        all_folders[role.name] = {
            "accepted": os.path.join(base, "accepted"),
            "doubtful": os.path.join(base, "doubtful"),
            "rejected": os.path.join(base, "rejected"),
        }
        for folder in all_folders[role.name].values():
            os.makedirs(folder, exist_ok=True)

    # "Other" bucket
    other_base = "other"
    other_folders = {
        "accepted": os.path.join(other_base, "accepted"),
        "doubtful": os.path.join(other_base, "doubtful"),
        "rejected": os.path.join(other_base, "rejected"),
    }
    for folder in other_folders.values():
        os.makedirs(folder, exist_ok=True)

    if not os.path.exists("cvs"):
        raise FileNotFoundError("'cvs' folder not found.")

    # Collect all PDF files
    pdf_files = []
    for root, _, files in os.walk("cvs"):
        for f in files:
            if f.lower().endswith(".pdf"):
                pdf_files.append((os.path.join(root, f), f))

    total_cvs = len(pdf_files)
    current_iteration = 0

    # Statistics: per role (including "Other")
    stats = {role.name: {"accepted": 0, "doubtful": 0, "rejected": 0, "total": 0} for role in roles}
    stats["Other"] = {"accepted": 0, "doubtful": 0, "rejected": 0, "total": 0}

    for cv_idx, (path, filename) in enumerate(pdf_files):
        current_iteration += 1
        progress = {
            "current": current_iteration,
            "total": total_cvs,
            "cv": f"{cv_idx+1}/{total_cvs}"
        }

        print(f"\n📄 [{cv_idx+1}/{total_cvs}] {filename}")

        text = extract_pdf_text(path)
        if not text.strip():
            assigned_role = "Other"
            decision = "Rejected"
            reason = "Empty or unreadable PDF"
            certainty_score = 0
            extracted_degree = ""
            extracted_qualification = ""
            extracted_experience = 0.0
            shutil.copy(path, os.path.join(other_folders["rejected"], filename))
            stats["Other"]["rejected"] += 1
            stats["Other"]["total"] += 1
        else:
            best_score = -1
            best_role = None
            best_result = None

            for role in roles:
                print(f"   🤖 Evaluating against role: {role.name}")
                result = evaluate_cv(text, role)
                score = result.get("certainty_score", 0)
                print(f"      Score: {score}")
                if score > best_score:
                    best_score = score
                    best_role = role.name
                    best_result = result

            # ----- CHANGED: ASSIGNMENT THRESHOLD RAISED TO 40 -----
            if best_score >= 40:
                assigned_role = best_role
                result = best_result
                extracted_degree = result.get("extracted_degree", "")
                extracted_qualification = result.get("extracted_qualification", "")
                extracted_experience = result.get("extracted_experience", 0.0)
                certainty_score = result.get("certainty_score", 0)
                reason = result.get("reason", "No reason provided")

                # ----- DECISION BOUNDARIES UNCHANGED -----
                if certainty_score >= 80:
                    decision = "Accepted"
                elif 50 <= certainty_score < 80:
                    decision = "Doubtful"
                else:
                    decision = "Rejected"

                dest_folder = all_folders[assigned_role][decision.lower()]
                shutil.copy(path, os.path.join(dest_folder, filename))
                stats[assigned_role][decision.lower()] += 1
                stats[assigned_role]["total"] += 1
            else:
                assigned_role = "Other"
                decision = "Rejected"
                reason = "Does not meet minimum threshold (40) for any role"
                certainty_score = best_score
                extracted_degree = ""
                extracted_qualification = ""
                extracted_experience = 0.0
                shutil.copy(path, os.path.join(other_folders["rejected"], filename))
                stats["Other"]["rejected"] += 1
                stats["Other"]["total"] += 1

        result_event = {
            "type": "result",
            "role": assigned_role,
            "data": {
                "filename": filename,
                "decision": decision,
                "reason": reason,
                "certainty_score": certainty_score,
                "extracted_degree": extracted_degree,
                "extracted_qualification": extracted_qualification,
                "extracted_experience": extracted_experience
            },
            "progress": progress
        }
        yield result_event

        if (cv_idx + 1) % 5 == 0 and cv_idx + 1 < total_cvs:
            print(f"⏳ [COOL-OFF] {cv_idx+1}/{total_cvs} CVs done, sleeping 30s...")
            time.sleep(30)
            print("🟢 Resuming...\n")

    # Build summary (includes "Other")
    summary_data = {}
    for role_name, s in stats.items():
        summary_data[role_name] = {
            "total": s["total"],
            "accepted": s["accepted"],
            "doubtful": s["doubtful"],
            "rejected": s["rejected"]
        }
    yield {"type": "summary", "data": summary_data}